#!/usr/bin/env python3
"""Update AgentCore_Technical_Flow.docx with sequence diagram images.

Appends a new "Sequence Diagrams" section at the end of the existing .docx file,
embedding the generated PNG diagrams with explanatory text.

Requirements:
    pip install python-docx

Usage:
    python docs/update_technical_doc.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(SCRIPT_DIR, "AgentCore_Technical_Flow.docx")

DIAGRAM_1_PATH = os.path.join(SCRIPT_DIR, "sequence_diagram_access_control.png")
DIAGRAM_2_PATH = os.path.join(SCRIPT_DIR, "sequence_diagram_personas.png")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    """Add a heading paragraph to the document."""
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False, font_size=None):
    """Add a simple paragraph with optional formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    return para


def add_image(doc, image_path, width_inches=6.5, caption=None):
    """Add an image centered with an optional caption below."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    # Verify files exist
    if not os.path.isfile(DOCX_PATH):
        print(f"ERROR: Document not found: {DOCX_PATH}")
        return

    for path in [DIAGRAM_1_PATH, DIAGRAM_2_PATH]:
        if not os.path.isfile(path):
            print(f"ERROR: Diagram not found: {path}")
            print("       Run create_sequence_diagram.py first to generate diagrams.")
            return

    # Open existing document
    doc = Document(DOCX_PATH)

    # Add page break before new section
    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Section: Sequence Diagrams
    # -----------------------------------------------------------------------
    add_heading(doc, "Sequence Diagrams: Access Control Flow", level=1)

    add_paragraph(
        doc,
        "The following diagrams illustrate the end-to-end access control flow "
        "implemented by AgentCore Gateway. They show how user identity (via JWT "
        "claims from Cognito) flows through the REQUEST Interceptor, gets evaluated "
        "by the Cedar Policy Engine, and determines whether an MCP tool is invoked "
        "or the request is denied — all before any data leaves the system.",
    )

    # Diagram 1: Full sequence diagram
    add_heading(doc, "Diagram 1: End-to-End Access Control Sequence", level=2)

    add_paragraph(
        doc,
        "This UML-style sequence diagram shows the complete lifecycle of a request, "
        "from user authentication through Cognito, to the agent's LLM reasoning, "
        "through the AgentCore Gateway's enrichment and authorization layers, and "
        "finally to tool execution or denial.",
    )

    add_paragraph(doc, "Key observations:", bold=True)

    bullets = [
        "Authentication: Cognito issues a JWT with custom claims (role, line_scope, "
        "plant_scope, equipment_scope) that encode the user's data boundaries.",
        "Enrichment: The REQUEST Interceptor decodes the JWT and injects user_context "
        "into the tool call arguments — the agent/LLM never sees or controls this step.",
        "Authorization: Cedar evaluates the enriched request against policies. "
        "A 'permit_all' baseline allows access, while 'forbid_*' rules check whether "
        "the requested parameter falls within the user's scope.",
        "Execution vs Denial: If Cedar allows, the Gateway forwards to the Lambda tool "
        "target. If Cedar denies, the MCP server is NEVER invoked — zero data exposure.",
        "Synthesis: The agent receives either tool results or a deny message, and "
        "composes a natural-language response. On denial, it explains the boundary and "
        "suggests alternatives within scope.",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    add_image(
        doc,
        DIAGRAM_1_PATH,
        width_inches=6.5,
        caption="Figure 1: AgentCore Gateway — Fine-Grained Access Control Sequence Diagram",
    )

    # Diagram 2: Three personas comparison
    add_heading(doc, "Diagram 2: Three Personas — Same Interface, Different Access", level=2)

    add_paragraph(
        doc,
        "This comparison diagram demonstrates how the same agent interface produces "
        "different outcomes for different users. Each row represents a persona with "
        "a horizontal left-to-right flow through the system components.",
    )

    add_paragraph(doc, "Personas:", bold=True)

    personas_desc = [
        "Sarah Chen (Plant Manager) — Full access. Queries are never denied because "
        "her role has the 'has_full_access' attribute, bypassing all forbid rules.",
        "Raj Patel (Line Supervisor) — Scoped to Line 7, Plant 2. Attempts to access "
        "Line 4 data are denied at the Cedar layer. The Lambda tool is never invoked.",
        "Priya Nair (Maintenance Technician) — Scoped to Machines 41-45. Attempts to "
        "access Machine 72 are denied. The agent suggests checking Machine 42 instead.",
    ]
    for desc in personas_desc:
        doc.add_paragraph(desc, style="List Bullet")

    add_paragraph(doc, "AgentCore components highlighted:", bold=True)
    add_paragraph(
        doc,
        "The purple-shaded region in the diagram encloses the two AgentCore-managed "
        "components: the Gateway (with its REQUEST Interceptor) and the Cedar Policy "
        "Engine. These operate server-side, external to the agent process, ensuring "
        "that neither the LLM nor application code can bypass authorization.",
    )

    add_image(
        doc,
        DIAGRAM_2_PATH,
        width_inches=6.5,
        caption="Figure 2: Same Agent, Same Interface — Different Data Access "
                "(AgentCore Gateway enforces Cedar policies at the parameter level)",
    )

    # -----------------------------------------------------------------------
    # Section: Security Properties (summary)
    # -----------------------------------------------------------------------
    add_heading(doc, "Key Security Properties Illustrated", level=2)

    properties = [
        ("Deterministic Authorization", 
         "Cedar evaluation is pure logic — same input always produces the same output. "
         "The LLM cannot influence the policy decision regardless of prompt injection."),
        ("Fail-Secure", 
         "If the interceptor crashes, no context is injected. Cedar has nothing to "
         "permit, so the result is DENY (deny-by-default)."),
        ("MCP Server Isolation", 
         "Denied requests NEVER reach the MCP server. The tool Lambda is not invoked. "
         "No data processing occurs. No compute is wasted."),
        ("Graceful Degradation", 
         "The agent receives the deny message as a tool result and intelligently "
         "responds: explains what it CAN access, suggests alternatives, and does NOT "
         "retry the denied call."),
    ]

    for title, description in properties:
        para = doc.add_paragraph()
        run_title = para.add_run(f"{title}: ")
        run_title.bold = True
        para.add_run(description)

    # Save
    doc.save(DOCX_PATH)
    print(f"SUCCESS: Updated {DOCX_PATH}")
    print(f"  - Added 'Sequence Diagrams: Access Control Flow' section")
    print(f"  - Embedded: {os.path.basename(DIAGRAM_1_PATH)}")
    print(f"  - Embedded: {os.path.basename(DIAGRAM_2_PATH)}")


if __name__ == "__main__":
    main()
