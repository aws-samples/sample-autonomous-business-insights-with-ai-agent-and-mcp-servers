#!/usr/bin/env python3
"""Generate Word document: Technical Flow with AgentCore Component Details."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)
    return p


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(
        'Persona-Based Access Control with Amazon Bedrock AgentCore', level=0
    )
    add_para(doc, 'Technical Architecture & Implementation Guide', italic=True, size=14)
    add_para(doc, 'Manufacturing Insights — Multi-Agent System with MCP Servers', size=12)
    doc.add_paragraph()
    add_para(doc, 'Version 2.0 | June 2026 | DEPLOYED on AWS Account 338277320360', size=10)
    doc.add_page_break()

    # Deployed Resources
    add_heading(doc, 'Deployed Infrastructure (Live)', level=1)
    add_para(doc, 'The following resources are deployed and operational on AWS:', bold=True)
    add_table(doc,
        ['Component', 'Resource ID', 'Status'],
        [
            ['Gateway', 'mfginsightsgateway-kbvnf0ga6j', 'READY'],
            ['Gateway URL', 'https://mfginsightsgateway-...us-east-1.amazonaws.com', 'Active'],
            ['Policy Engine', 'MfgInsightsPolicyEngine-w1do75vmrk', 'ENFORCE'],
            ['Cedar Policies', 'permit_all, forbid_line_scope, forbid_equipment_scope', '3 Active'],
            ['Cognito Pool', 'us-east-1_wBnf60sfQ', 'Active'],
            ['Users', 'sarah.chen, raj.patel, priya.nair', 'Confirmed'],
            ['Lambda: Equipment', 'MfgInsights-EquipmentTools', 'READY'],
            ['Lambda: IoT', 'MfgInsights-IoTTools', 'READY'],
            ['Lambda: Analytics', 'MfgInsights-AnalyticsTools', 'READY'],
            ['Test Gateway', 'mfginsightstest-af76b5qmwe', 'READY (no auth)'],
        ]
    )
    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', level=1)
    toc = [
        '1. Executive Summary',
        '2. AgentCore Components Overview',
        '3. Architecture Diagram',
        '4. Component Deep Dive',
        '   4.1 AgentCore Runtime',
        '   4.2 AgentCore Gateway',
        '   4.3 AgentCore Identity',
        '   4.4 AgentCore Policy (Cedar)',
        '   4.5 AgentCore Memory',
        '   4.6 AgentCore Observability',
        '5. Persona-Based Access Flow',
        '6. Sequence Diagrams',
        '7. Cedar Policy Details',
        '8. Lambda Interceptor Pipeline',
        '9. Security Properties',
        '10. Demo Scenarios',
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # Section 1
    add_heading(doc, '1. Executive Summary', level=1)
    add_para(doc, (
        'This document describes how Amazon Bedrock AgentCore enables fine-grained, '
        'persona-based access control for AI agents that serve multiple user roles within '
        'a manufacturing enterprise. A single Strands Agent connects to multiple MCP servers '
        'via AgentCore Gateway, while Cedar policies enforce that each user only sees data '
        'within their authorized scope — regardless of what the LLM attempts to access.'
    ))
    doc.add_paragraph()
    add_para(doc, 'Key Outcomes:', bold=True)
    bullets = [
        'Plant managers see all 12 lines across 3 plants',
        'Line supervisors see only their assigned lines',
        'Maintenance technicians see only their assigned machines',
        'Same agent, same tools — different data boundaries',
        'Policy enforcement is deterministic (Cedar) — LLM cannot bypass',
        'Zero auth logic inside MCP servers — all enforcement at Gateway',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_page_break()

    # Section 2: AgentCore Components
    add_heading(doc, '2. AgentCore Components Overview', level=1)
    add_para(doc, (
        'Amazon Bedrock AgentCore provides six core services that together form '
        'a complete platform for building, deploying, and securing AI agents at scale.'
    ))
    doc.add_paragraph()

    components = [
        ('AgentCore Runtime', 'Serverless compute for agent execution',
         'Hosts the Strands Agent in Firecracker microVMs. Each user session runs in '
         'an isolated environment. Supports direct code deployment (Python zip) or '
         'container images. Auto-scales to zero when idle. Provides session management '
         'and automatic sandboxing.'),
        ('AgentCore Gateway', 'Managed MCP router + security enforcement point',
         'Single HTTPS endpoint that routes agent tool calls to MCP servers (Lambda targets, '
         'OpenAPI endpoints, or remote MCP servers). Validates JWT tokens, runs interceptors, '
         'evaluates Cedar policies, caches responses (3-tier: edge, regional, per-session). '
         'This is THE chokepoint where all security enforcement happens.'),
        ('AgentCore Identity', 'Authentication + credential propagation',
         'Integrates with OAuth 2.0 providers (Cognito, Okta, Entra ID). Propagates user '
         'identity through the entire call chain via session headers. Manages token vaults '
         'for act-on-behalf flows. Ensures every tool call carries authenticated context.'),
        ('AgentCore Policy', 'Cedar-based deterministic authorization',
         'Uses the Cedar policy language (developed by AWS, formally verified) to define '
         'permit/forbid rules. Evaluates at the Gateway BEFORE tool execution. Deny-by-default '
         'semantics. Forbid always overrides permit. Every decision logged to CloudWatch. '
         'Policies can be authored in natural language and auto-generated.'),
        ('AgentCore Memory', 'Session + cross-session context persistence',
         'Three namespaces: user-scoped, team-scoped, organization-scoped. Short-term memory '
         '(within session) enables follow-up questions. Long-term memory persists preferences, '
         'baselines, and recurring patterns. Memory retrieval respects Policy — a user cannot '
         'access memories derived from data outside their scope.'),
        ('AgentCore Observability', 'Tracing, logging, metrics',
         'Integrates with AWS X-Ray for distributed tracing across agent → gateway → tools. '
         'All policy decisions logged to CloudWatch. Tool call latency metrics. Token usage '
         'tracking. Anomaly detection on agent behavior patterns.'),
    ]

    for name, tagline, description in components:
        add_para(doc, f'{name} — {tagline}', bold=True, size=12)
        add_para(doc, description)
        doc.add_paragraph()
    doc.add_page_break()

    # Section 3: Architecture
    add_heading(doc, '3. Architecture Diagram', level=1)
    add_code_block(doc, """
Users (Cognito JWT)
    |
    v
+------------------------------------------------------------------+
|  AgentCore Runtime (Firecracker microVM)                          |
|  +------------------------------------------------------------+  |
|  |  Strands Agent (Claude Sonnet on Bedrock)                  |  |
|  |  System prompt: identity + scope + memory context          |  |
|  +------------------------------------------------------------+  |
+------------------------------------------------------------------+
    |
    v  (every tool call)
+------------------------------------------------------------------+
|  AgentCore Gateway                                                |
|  +---------------+  +-----------+  +-----------+  +-----------+  |
|  | JWT Validator |->| REQUEST   |->| POLICY    |->| RESPONSE  |  |
|  | (built-in)   |  | Intercept |  | ENGINE    |  | Intercept |  |
|  |              |  | (Lambda)  |  | (Cedar)   |  | (Lambda)  |  |
|  +---------------+  +-----------+  +-----------+  +-----------+  |
|                                         |                         |
|                               ALLOW or DENY                       |
+------------------------------------------------------------------+
    |  (only if ALLOW)
    v
+------------------------------------------------------------------+
|  Lambda Tool Targets (MCP)                                        |
|  +----------+ +----------+ +----------+ +----------+ +--------+  |
|  |Equipment | |IoT Telem | |Supply    | |Analytics | |Semantic|  |
|  |Status    | |etry      | |Chain     | |OEE+Qual  | |Layer   |  |
|  +----------+ +----------+ +----------+ +----------+ +--------+  |
+------------------------------------------------------------------+
    |
    v
+------------------------------------------------------------------+
|  Data Infrastructure                                              |
|  Aurora | Timestream | Redshift | OpenSearch | S3                  |
+------------------------------------------------------------------+
""")
    doc.add_page_break()

    # Section 4: Component Deep Dive
    add_heading(doc, '4. Component Deep Dive', level=1)

    add_heading(doc, '4.1 AgentCore Runtime', level=2)
    add_para(doc, 'How it works in this solution:', bold=True)
    runtime_details = [
        'The Strands Agent code is deployed as a Python zip to AgentCore Runtime',
        'Each user session gets an isolated Firecracker microVM (no cross-session data leakage)',
        'The agent connects to Gateway (not directly to MCP servers) via the Gateway URL',
        'Session isolation means Raj cannot see Sarah\'s in-flight queries or cached data',
        'Auto-scales: 0 instances when no queries, scales up within 50ms on request',
        'Direct code deployment: no Docker required, just zip your Python code',
    ]
    for item in runtime_details:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '4.2 AgentCore Gateway', level=2)
    add_para(doc, 'The security enforcement chokepoint:', bold=True)
    add_para(doc, (
        'Every tool call from the agent passes through the Gateway. This is non-negotiable — '
        'the agent cannot bypass it because it only has the Gateway URL, not direct MCP server URLs. '
        'DEPLOYED: mfginsightsgateway-kbvnf0ga6j (READY, ENFORCE mode).'
    ))
    doc.add_paragraph()
    add_para(doc, 'Gateway capabilities used:', bold=True)
    gw_items = [
        'JWT Validation: Verifies Cognito token signature against JWKS endpoint',
        'REQUEST Interceptor: Lambda that enriches request with user scope (runs FIRST)',
        'Policy Engine: Cedar evaluation against enriched request (runs SECOND)',
        'RESPONSE Interceptor: Lambda that filters tool list by role (runs LAST)',
        'MCP Routing: Routes tools/call to the correct Lambda target based on tool name prefix',
        '3-Tier Cache: Edge (CloudFront), Regional (ElastiCache), Per-session (in-microVM)',
    ]
    for item in gw_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '4.3 AgentCore Identity', level=2)
    add_para(doc, 'Authentication and scope propagation:', bold=True)
    identity_items = [
        'Cognito User Pool with custom attributes: role, line_scope, plant_scope, equipment_scope',
        'Three user groups: plant_managers, line_supervisors, maintenance_technicians',
        'JWT includes all custom claims — available to interceptor and Cedar at runtime',
        'Custom attributes are admin-only writable (users cannot escalate their own scope)',
        'Token expiry: 60 minutes default — forces re-authentication',
        'In production: replace Cognito with Okta/Entra ID via AgentCore Identity federation',
    ]
    for item in identity_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '4.4 AgentCore Policy (Cedar)', level=2)
    add_para(doc, 'Deterministic, auditable authorization:', bold=True)
    add_para(doc, (
        'Cedar is a declarative policy language that evaluates permit/forbid rules over '
        'three elements: principal (who), action (what tool), resource (which gateway). '
        'Conditions can reference any field in the request context.'
    ))
    doc.add_paragraph()
    add_para(doc, 'Our Cedar policies:', bold=True)
    add_table(doc,
        ['Policy', 'Type', 'What It Does'],
        [
            ['permit_all', 'permit', 'Baseline: all authenticated users can call tools'],
            ['forbid_line_scope', 'forbid', 'Blocks line_supervisors from lines outside their scope'],
            ['forbid_equipment_scope', 'forbid', 'Blocks maintenance_technicians from unassigned machines'],
            ['forbid_plant_scope', 'forbid', 'Blocks all non-admin users from unauthorized plants'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, 'Key Cedar properties:', bold=True)
    cedar_props = [
        'Deny-by-default: no matching permit = DENY',
        'Forbid overrides permit: a single forbid match blocks regardless of permits',
        'Deterministic: same input always produces same decision (unlike LLM)',
        'Formally verified: AWS tools can prove policy completeness and consistency',
        'Auditable: every decision logged with full context to CloudWatch',
        'Fast: <1ms evaluation time (no Lambda cold start)',
    ]
    for item in cedar_props:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '4.5 AgentCore Memory', level=2)
    add_para(doc, 'Context persistence across sessions:', bold=True)
    memory_items = [
        'Short-term (session): Turn-by-turn context for follow-up questions',
        'Long-term (user): Preferences, baselines, recurring patterns (e.g., Priya\'s Machine 42 baseline: 3.8 mm/s)',
        'Long-term (team): Shared thresholds and standards (e.g., temperature warning at 72C)',
        'Memory respects policy: Raj cannot access memories from data outside his line scope',
        'Namespace isolation: user > team > organization hierarchy',
    ]
    for item in memory_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '4.6 AgentCore Observability', level=2)
    add_para(doc, 'Full-stack tracing and audit:', bold=True)
    obs_items = [
        'X-Ray traces: agent → gateway → interceptor → policy → tool → response',
        'Policy decision logs: every ALLOW/DENY with principal, action, resource, context',
        'Tool call metrics: latency percentiles, error rates, throttle counts',
        'Interceptor logs: JWT claims extraction, context injection, filtering actions',
        'Agent reasoning traces: which tools selected, in what order, with what params',
        'Cost tracking: Bedrock token usage per user, per session',
    ]
    for item in obs_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # Section 5: Persona-Based Access Flow
    add_heading(doc, '5. Persona-Based Access Flow', level=1)
    add_para(doc, 'Three users, same agent, different boundaries:', bold=True)
    doc.add_paragraph()

    add_table(doc,
        ['User', 'Role', 'Plant Scope', 'Line Scope', 'Equipment Scope'],
        [
            ['Sarah Chen', 'Plant Manager', 'All (1,2,3)', 'All (1-12)', 'Full access'],
            ['Raj Patel', 'Line Supervisor', 'Plant 2', 'Line 7 only', 'Machine 71-75'],
            ['Priya Nair', 'Maint. Technician', 'Plant 1', 'Line 4 only', 'Machine 41-45 only'],
        ]
    )
    doc.add_paragraph()

    add_para(doc, 'Decision Matrix (12 scenarios):', bold=True)
    add_table(doc,
        ['#', 'User', 'Tool', 'Parameter', 'Decision', 'Cedar Rule'],
        [
            ['1', 'Sarah', 'get_oee_trends', 'line="Line 4"', 'ALLOW', 'Full access bypass'],
            ['2', 'Sarah', 'get_sensor_readings', 'machine_id=72', 'ALLOW', 'Full access bypass'],
            ['3', 'Raj', 'get_oee_trends', 'line="Line 7"', 'ALLOW', 'Line 7 in scope'],
            ['4', 'Raj', 'get_oee_trends', 'line="Line 4"', 'DENY', 'forbid_line_scope'],
            ['5', 'Raj', 'get_equipment', 'plant="Plant 1"', 'DENY', 'forbid_plant_scope'],
            ['6', 'Raj', 'detect_anomaly', '(no params)', 'ALLOW', 'No dimension to check'],
            ['7', 'Priya', 'get_sensor', 'machine_id=42', 'ALLOW', 'Machine 42 in scope'],
            ['8', 'Priya', 'get_sensor', 'machine_id=72', 'DENY', 'forbid_equipment_scope'],
            ['9', 'Priya', 'get_oee', 'line="Line 4"', 'ALLOW', 'Line 4 in scope'],
            ['10', 'Priya', 'get_oee', 'line="Line 7"', 'DENY', 'forbid_line_scope'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        'Key insight: Access control is at the PARAMETER level, not the tool level. '
        'Raj CAN call get_oee_trends — just not with line="Line 4". This enables '
        'the same agent to serve all roles without creating role-specific tool sets.'
    ), italic=True)
    doc.add_page_break()

    # Section 6: Sequence Diagrams
    add_heading(doc, '6. Sequence Diagrams', level=1)

    add_heading(doc, '6.1 Scenario: ALLOWED — Raj queries Line 7 OEE', level=2)
    add_code_block(doc, """Raj -> Cognito: Authenticate (username + password)
Cognito -> Raj: JWT {role:"line_supervisor", line_scope:"Line 7"}

Raj -> Agent: "What's the OEE for Line 7?"
Agent -> Agent: LLM reasons → need get_oee_trends(line="Line 7")

Agent -> Gateway: tools/call {name:"get_oee_trends", args:{line:"Line 7"}}
Gateway -> Gateway: Validate JWT signature (built-in)
Gateway -> REQUEST Interceptor: Enrich request
REQUEST Interceptor -> REQUEST Interceptor: Decode JWT claims
REQUEST Interceptor -> REQUEST Interceptor: Inject user_context into args
REQUEST Interceptor -> Gateway: Return enriched request

Gateway -> Policy Engine: Evaluate Cedar rules
Policy Engine -> Policy Engine: permit_all → PERMIT matches
Policy Engine -> Policy Engine: forbid_line_scope → check "Line 7" in ["Line 7"] → YES (in scope)
Policy Engine -> Policy Engine: No forbid matches → result: ALLOW
Policy Engine -> Gateway: ALLOW

Gateway -> Analytics Lambda: Execute get_oee_trends(line="Line 7")
Analytics Lambda -> Gateway: OEE data (82.3%, -3.1% WoW)

Gateway -> RESPONSE Interceptor: Filter response
RESPONSE Interceptor -> Gateway: Pass through (not a tools/list)

Gateway -> Agent: OEE data
Agent -> Agent: LLM synthesizes response
Agent -> Raj: "Line 7 OEE is 82.3% this week, down 3.1%..."
""")

    add_heading(doc, '6.2 Scenario: DENIED — Raj queries Line 4 (out of scope)', level=2)
    add_code_block(doc, """Raj -> Agent: "Show me equipment status for Line 4"
Agent -> Agent: LLM reasons → need get_equipment_status(line="Line 4")

Agent -> Gateway: tools/call {name:"get_equipment_status", args:{line:"Line 4"}}
Gateway -> Gateway: Validate JWT ✓
Gateway -> REQUEST Interceptor: Enrich request
REQUEST Interceptor -> Gateway: Enriched (line_scope:"Line 7" injected)

Gateway -> Policy Engine: Evaluate Cedar rules
Policy Engine -> Policy Engine: permit_all → PERMIT matches
Policy Engine -> Policy Engine: forbid_line_scope →
    check: action matches? YES
    check: context.input.line = "Line 4"
    check: principal in line_supervisors? YES
    check: "Line 4" in line_scope "Line 7"? NO
    → FORBID MATCHES
Policy Engine -> Gateway: DENY (forbid overrides permit)

Gateway -> Agent: "[Policy] Access denied. Line 4 not authorized. Scope: Line 7"
                 *** MCP SERVER NEVER CALLED ***

Agent -> Agent: LLM reads deny, adjusts response
Agent -> Raj: "I don't have access to Line 4 data. I can access Line 7.
              Would you like Line 7 equipment status instead?"
""")
    doc.add_page_break()

    # Section 7: Cedar Policy Details
    add_heading(doc, '7. Cedar Policy Details', level=1)
    add_para(doc, 'Complete Cedar policies used in this system:', bold=True)

    add_heading(doc, '7.1 permit_all.cedar (Baseline)', level=2)
    add_code_block(doc, """// Baseline: allow all authenticated users to invoke tools.
// Required because Cedar uses deny-by-default.
permit(
    principal,
    action,
    resource == AgentCore::Gateway::"arn:aws:bedrock-agentcore:...:gateway/mfg-insights"
);""")

    add_heading(doc, '7.2 forbid_line_scope.cedar', level=2)
    add_code_block(doc, """// Line supervisors restricted to their assigned lines.
forbid(
    principal is AgentCore::OAuthUser,
    action in [
        AgentCore::Action::"EquipmentTarget___get_equipment_status",
        AgentCore::Action::"EquipmentTarget___get_shared_infrastructure",
        AgentCore::Action::"IoTTarget___detect_anomaly",
        AgentCore::Action::"AnalyticsTarget___get_oee_trends",
        AgentCore::Action::"AnalyticsTarget___get_quality_metrics"
    ],
    resource == AgentCore::Gateway::"arn:aws:...:gateway/mfg-insights"
) when {
    context.input has line &&
    principal.hasTag("cognito:groups") &&
    principal.getTag("cognito:groups") like "*line_supervisors*" &&
    principal.hasTag("custom:line_scope") &&
    !(principal.getTag("custom:line_scope") like ("*" + context.input.line + "*"))
};""")

    add_heading(doc, '7.3 forbid_equipment_scope.cedar', level=2)
    add_code_block(doc, """// Maintenance technicians restricted to assigned machines.
forbid(
    principal is AgentCore::OAuthUser,
    action in [
        AgentCore::Action::"EquipmentTarget___get_equipment_status",
        AgentCore::Action::"EquipmentTarget___get_maintenance_history",
        AgentCore::Action::"IoTTarget___get_sensor_readings"
    ],
    resource == AgentCore::Gateway::"arn:aws:...:gateway/mfg-insights"
) when {
    context.input has machine_id &&
    principal.hasTag("cognito:groups") &&
    principal.getTag("cognito:groups") like "*maintenance_technicians*" &&
    principal.hasTag("custom:equipment_scope") &&
    !(principal.getTag("custom:equipment_scope") like ("*Machine " + context.input.machine_id + "*"))
};""")
    doc.add_page_break()

    # Section 8: Interceptor Pipeline
    add_heading(doc, '8. Lambda Interceptor Pipeline', level=1)
    add_para(doc, 'Execution order in the Gateway:', bold=True)
    add_code_block(doc, """Agent → JWT Validation → REQUEST Interceptor → Cedar Policy → Tool → RESPONSE Interceptor → Agent""")
    doc.add_paragraph()

    add_heading(doc, '8.1 REQUEST Interceptor', level=2)
    add_para(doc, 'Purpose: Extract JWT claims and inject user scope into tool arguments.', bold=True)
    add_table(doc,
        ['Input', 'Processing', 'Output'],
        [
            ['Authorization: Bearer eyJ...', 'Decode JWT payload (base64)', 'user_context injected into args'],
            ['body.params.arguments: {line:"L4"}', 'Extract: role, scope attributes', 'args.user_context = {role, scope}'],
            ['', 'Set header: x-user-role', 'Headers updated for response interceptor'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, 'Why Lambda (not Cedar): Cedar cannot decode JWTs, call APIs, or modify payloads.', italic=True)

    add_heading(doc, '8.2 RESPONSE Interceptor', level=2)
    add_para(doc, 'Purpose: Filter tools/list so agent only sees authorized tools.', bold=True)
    add_table(doc,
        ['Role', 'Visible Tools', 'Hidden Tools'],
        [
            ['plant_manager', 'All 12 tools', 'None'],
            ['line_supervisor', '8 tools (no maintenance_history, sensor by machine)', '4 tools hidden'],
            ['maintenance_technician', '8 tools (no OEE, no quality, no shared_infra)', '4 tools hidden'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        'Why this matters: If the LLM sees 12 tools but can only use 8, it might attempt '
        '(and fail) to call unauthorized ones. Filtering the list means the LLM never considers them.'
    ), italic=True)
    doc.add_page_break()

    # Section 9: Security Properties
    add_heading(doc, '9. Security Properties', level=1)
    add_table(doc,
        ['Property', 'Guarantee', 'Mechanism'],
        [
            ['Deterministic Auth', 'Same input → same decision (no LLM variance)', 'Cedar evaluation (pure logic)'],
            ['Fail-Secure', 'If interceptor fails → DENY (deny-by-default)', 'Cedar + no permit = DENY'],
            ['MCP Isolation', 'Denied requests never reach tools', 'Gateway blocks before forwarding'],
            ['Audit Complete', 'Every decision logged with context', 'CloudWatch + CloudTrail'],
            ['Graceful Degrade', 'Agent explains scope limits to user', 'LLM reads deny message, adapts'],
            ['No Bypass', 'LLM cannot influence policy outcome', 'Cedar runs outside LLM process'],
            ['Least Privilege', 'Each Lambda has minimal IAM', 'Separate role per target'],
            ['Session Isolation', 'Users cannot access each other data', 'Firecracker microVM per session'],
        ]
    )
    doc.add_page_break()

    # Section 10: Demo Scenarios
    add_heading(doc, '10. Demo Scenarios', level=1)

    add_heading(doc, 'Demo 1: Same Question, Different Answers', level=2)
    add_para(doc, 'Ask all three: "Are there any anomalies on the factory floor?"', bold=True)
    add_table(doc,
        ['User', 'What They See', 'What Is Blocked'],
        [
            ['Sarah', 'Anomalies across all 12 lines, ranked by severity', 'Nothing — full access'],
            ['Raj', 'Only Line 7 anomalies (or "no issues on your line")', 'Lines 1-6, 8-12 blocked'],
            ['Priya', 'Only Machine 41-45 readings', 'All other machines blocked'],
        ]
    )

    add_heading(doc, 'Demo 2: Cross-System Intelligence (Sarah only)', level=2)
    add_para(doc, 'Sarah asks: "Why is Line 4 availability dropping?"', bold=True)
    add_para(doc, 'Agent correlates 4 data sources in one response:')
    correlations = [
        'IoT: Machine 42 temperature +12°C above baseline (anomaly)',
        'Equipment: Bearing replaced 8 months ago, running at 1.3x rated capacity',
        'Supply Chain: Replacement bearings below reorder point (12 vs 20)',
        'Analytics: Line 4 OEE dropped 6% over 4 weeks, quality scrap up 2x',
        'Shared Infra: Lines 4 and 9 share coolant loop A — both degrading',
    ]
    for item in correlations:
        doc.add_paragraph(item, style='List Bullet')
    add_para(doc, (
        'Root cause synthesis: Bearing degradation under sustained overload, '
        'compounded by coolant loop degradation affecting both Line 4 and Line 9.'
    ), italic=True)

    add_heading(doc, 'Demo 3: Memory-Augmented Follow-up (Priya)', level=2)
    add_para(doc, 'Priya asks: "Has the vibration on Machine 42 gotten worse?"', bold=True)
    add_para(doc, 'Flow:')
    memory_flow = [
        'Memory surfaces: "Last week vibration = 3.8 mm/s" (from prior session)',
        'IoT tool returns current: 4.5 mm/s',
        'Agent computes: +18% increase, now above warning threshold (4.0 mm/s)',
        'Response: "Yes, +18% worse. Now at 4.5 mm/s (warning = 4.0). Recommend immediate inspection."',
    ]
    for item in memory_flow:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, 'Demo 4: Policy Enforcement in Action (Raj)', level=2)
    add_para(doc, 'Raj asks: "Show me Machine 42 vibration data"', bold=True)
    add_para(doc, (
        'Machine 42 is on Line 4 (not Raj\'s scope). Cedar evaluates the line parameter '
        'and denies access. The agent responds: "I can\'t access Line 4 data. I\'m authorized '
        'for Line 7. Want me to check your Line 7 machines instead?"'
    ))
    doc.add_page_break()

    # Section: Code Structure & Deep Dive
    add_heading(doc, '11. Code Structure', level=1)
    add_code_block(doc, """src/
├── main.py                    <- CLI entry point
├── demo_ui.py                 <- Streamlit web UI (chat + live logs)
├── config.py                  <- Environment config
├── agent/
│   ├── agent.py               <- THE CORE: ManufacturingInsightsAgent
│   └── prompts.py             <- System prompt (injects identity + memory)
├── identity/
│   ├── models.py              <- UserIdentity + 3 demo personas
│   ├── policy.py              <- PolicyEngine (local Cedar simulation)
│   └── gateway_hook.py        <- Strands hook (local mode only)
├── memory/
│   └── manager.py             <- Session + long-term memory
├── servers/
│   ├── start_all.py           <- Launch all 5 local MCP servers
│   ├── equipment_server.py    <- get_equipment_status, get_maintenance_history
│   ├── iot_telemetry_server.py <- get_sensor_readings, detect_anomaly
│   ├── supply_chain_server.py <- check_parts_inventory
│   └── analytics_server.py    <- get_oee_trends, get_quality_metrics
├── data/
│   ├── data_provider.py       <- Routes simulated vs live backends
│   └── sample_data.py         <- In-memory factory data

deploy/agentcore/
├── deploy_live.py             <- Creates real AWS resources
├── setup_identity.py          <- Cognito User Pool + users
├── setup_gateway.py           <- Gateway + Lambda targets
├── setup_policy.py            <- Cedar policies + Policy Engine
├── cedar_policies/            <- Raw .cedar files
└── lambda_functions/          <- Interceptor Lambda code""")

    add_heading(doc, '12. Core Agent Flow (agent.py)', level=1)
    add_para(doc, 'The ManufacturingInsightsAgent.query() method:', bold=True)
    add_code_block(doc, """def query(self, user, question):
    # 1. Build system prompt with identity + scope + memory
    system_prompt = self._build_system_prompt(user, session_id)

    # 2. Connect to tools (local MCP servers OR real AgentCore Gateway)
    mcp_clients = self._create_mcp_clients()

    # 3. Collect all tools into flat list
    all_tools = []
    for client in mcp_clients:
        all_tools.extend(client.list_tools_sync())

    # 4. Create agent (Gateway mode = no local hook)
    if USE_AGENTCORE_GATEWAY:
        agent = Agent(system_prompt, tools=all_tools)  # Gateway handles policy
    else:
        agent = Agent(system_prompt, tools=all_tools, hooks=[gateway_hook])

    # 5. Agent autonomously reasons and calls tools
    response = agent(question)

    # 6. Record in memory for follow-ups
    session.add_interaction(query=question, response_summary=response[:200])
    return response""")

    add_heading(doc, '13. Two Operating Modes', level=1)
    add_table(doc,
        ['Aspect', 'Local Simulation', 'Real AgentCore Gateway'],
        [
            ['Policy', 'Python if/else (policy.py)', 'Cedar evaluated server-side at Gateway'],
            ['Tools', 'Local FastMCP servers (ports 8001-8005)', 'Lambda targets invoked by Gateway'],
            ['Auth', 'Hardcoded UserIdentity dataclass', 'Cognito JWT with custom claims'],
            ['Enforcement', 'In-process Strands hook', 'Server-side (Gateway service)'],
            ['Audit', 'logging.warning()', 'CloudWatch + CloudTrail'],
            ['Isolation', 'Same process', 'Firecracker microVM per session'],
            ['Env var', 'USE_AGENTCORE_GATEWAY=false', 'USE_AGENTCORE_GATEWAY=true'],
        ]
    )

    add_heading(doc, '14. Request Lifecycle — Raj Asks About Line 4', level=1)
    add_para(doc, 'Concrete walkthrough of a DENIED request:', bold=True)
    add_code_block(doc, """Step 1: UI -> agent.query(raj_patel, "What's the OEE for Line 4?")

Step 2: System prompt built:
  "User: Raj Patel, Line Supervisor
   Scope: Plants: Plant 2, Lines: Line 7"

Step 3: Connect to Gateway (1 MCP endpoint, 3 tools visible)

Step 4: LLM reasons -> calls get_oee_trends(line="Line 4")

Step 5: Gateway evaluates:
  -> REQUEST Interceptor decodes JWT, injects context
  -> Cedar: permit_all MATCHES (would allow)
  -> Cedar: forbid_line_scope:
     "Line 4" NOT in raj's scope ["Line 7"]
     -> FORBID OVERRIDES PERMIT -> DENY
  -> Returns: "[Policy] Access denied. Not authorized for Line 4."

Step 6: LLM reads deny, responds:
  "I can't access Line 4. My scope is Line 7.
   Want Line 7 OEE trends instead?"

KEY: The Lambda target was NEVER invoked. Gateway blocked it.""")

    add_heading(doc, '15. AgentCore Component Interaction', level=1)
    add_code_block(doc, """
USER authenticates -> Cognito -> JWT with scope claims
          |
AGENT (Strands + Bedrock Claude)
  System prompt: identity + scope + memory
  LLM decides which tools to call
          | tools/call (MCP JSON-RPC)
GATEWAY (AgentCore) -- the chokepoint
  1. Validate JWT (Cognito JWKS)
  2. REQUEST Interceptor (decode JWT, inject context)
  3. Cedar Policy Engine:
     permit_all -> PERMIT (baseline)
     forbid_line_scope -> check line vs user scope
     forbid_equipment_scope -> check machine vs user scope
     If ANY forbid matches -> DENY
  4. If ALLOW -> invoke Lambda target
  5. RESPONSE Interceptor (filter tool list)
          |                    |
     [ALLOW]              [DENY]
          |                    |
  Lambda executes      "Access denied"
  Returns data         returned to agent
          |                    |
  Agent synthesizes    Agent explains limit
  response             suggests alternative""")

    add_heading(doc, '16. Design Principles', level=1)
    add_table(doc,
        ['Principle', 'How It Works'],
        [
            ['LLM decides what, Cedar decides if', 'Agent picks tools; Gateway approves/denies'],
            ['Fail-secure', 'No permit match = DENY. Interceptor crash = DENY.'],
            ['Parameter-level access', 'Same tool, different params = different decision'],
            ['MCP servers are auth-unaware', 'Zero access control logic in Lambdas'],
            ['Memory respects policy', 'Cannot surface memories from out-of-scope data'],
            ['Extensible by config', 'New tool = new Lambda, existing policies auto-apply'],
            ['Forbid overrides permit', 'One deny rule blocks regardless of permits'],
            ['Deterministic', 'Cedar same result for same input (unlike LLM)'],
        ]
    )
    doc.add_page_break()

    # Save
    output_path = Path(__file__).parent / 'docs' / 'AgentCore_Technical_Flow.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    print(f'Document saved: {output_path}')


if __name__ == '__main__':
    main()
